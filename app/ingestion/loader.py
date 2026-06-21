"""File loading / normalization for PDF, image, and DOCX inputs."""

from __future__ import annotations

import io
from pathlib import Path
from typing import Iterator

import fitz  # PyMuPDF
from PIL import Image, ImageOps

from app.config import get_settings
from app.ingestion.models import FileType, IngestedDocument, Page, UnsupportedFileError

# Extensions we accept, grouped by handler.
_PDF_EXT = {".pdf"}
_IMAGE_EXT = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff", ".gif"}
_DOCX_EXT = {".docx"}

# Magic-byte signatures for content sniffing (defends against wrong extensions).
_PNG_SIG = b"\x89PNG\r\n\x1a\n"
_JPG_SIG = b"\xff\xd8\xff"
_PDF_SIG = b"%PDF"
_ZIP_SIG = b"PK\x03\x04"  # docx/xlsx/zip


def detect_file_type(path: str | Path) -> FileType:
    """Determine the file type from extension, corroborated by magic bytes."""
    path = Path(path)
    ext = path.suffix.lower()

    head = b""
    try:
        with open(path, "rb") as fh:
            head = fh.read(8)
    except OSError:
        pass

    if head.startswith(_PDF_SIG) or ext in _PDF_EXT:
        return FileType.PDF
    if head.startswith(_PNG_SIG) or head.startswith(_JPG_SIG) or ext in _IMAGE_EXT:
        return FileType.IMAGE
    if ext in _DOCX_EXT and head.startswith(_ZIP_SIG):
        return FileType.DOCX
    if head.startswith(_ZIP_SIG) and ext in _DOCX_EXT:
        return FileType.DOCX
    return FileType.UNSUPPORTED


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
def _ingest_pdf(path: Path) -> IngestedDocument:
    settings = get_settings()
    try:
        doc = fitz.open(path)
    except Exception as exc:  # corrupt / not actually a PDF
        raise UnsupportedFileError(f"Could not open PDF '{path.name}': {exc}") from exc

    encrypted = bool(doc.needs_pass)
    if encrypted:
        # Try empty password; otherwise we can still rasterize? No — bail clearly.
        if not doc.authenticate(""):
            doc.close()
            raise UnsupportedFileError(f"PDF is password-protected: {path.name}")

    try:
        pages: list[Page] = []
        n = min(doc.page_count, settings.max_pages)
        for i in range(n):
            page = doc[i]
            text = page.get_text("text") or ""
            rect = page.rect
            pages.append(
                Page(
                    page_number=i + 1,
                    text=text.strip(),
                    width=int(rect.width),
                    height=int(rect.height),
                    has_embedded_text=len(text.strip()) > 0,
                )
            )
        meta = {"page_count_total": doc.page_count, "pages_ingested": n}
    except Exception as exc:  # damaged page tree etc.
        raise UnsupportedFileError(f"Could not read PDF '{path.name}': {exc}") from exc
    finally:
        doc.close()
    return IngestedDocument(
        filename=path.name,
        file_type=FileType.PDF,
        source_path=path,
        pages=pages,
        mime="application/pdf",
        encrypted=encrypted,
        metadata=meta,
    )


# --------------------------------------------------------------------------- #
# Image
# --------------------------------------------------------------------------- #
def _ingest_image(path: Path) -> IngestedDocument:
    try:
        with Image.open(path) as img:
            img = ImageOps.exif_transpose(img)  # honor camera rotation
            width, height = img.size
            fmt = (img.format or "").lower()
    except Exception as exc:  # corrupt / truncated image
        raise UnsupportedFileError(f"Could not open image '{path.name}': {exc}") from exc
    mime = f"image/{'jpeg' if fmt in ('jpg', 'jpeg') else (fmt or 'png')}"
    page = Page(page_number=1, text="", width=width, height=height, has_embedded_text=False)
    return IngestedDocument(
        filename=path.name,
        file_type=FileType.IMAGE,
        source_path=path,
        pages=[page],
        mime=mime,
        metadata={"image_format": fmt},
    )


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def _docx_text(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _ingest_docx(path: Path) -> IngestedDocument:
    text = _docx_text(path)
    page = Page(page_number=1, text=text, has_embedded_text=bool(text))
    return IngestedDocument(
        filename=path.name,
        file_type=FileType.DOCX,
        source_path=path,
        pages=[page],
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        metadata={"char_count": len(text)},
    )


# --------------------------------------------------------------------------- #
# Public API
# --------------------------------------------------------------------------- #
def ingest_file(path: str | Path) -> IngestedDocument:
    """Load a file into a normalized :class:`IngestedDocument`.

    Raises :class:`UnsupportedFileError` for types we cannot handle.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)

    ftype = detect_file_type(path)
    if ftype is FileType.PDF:
        return _ingest_pdf(path)
    if ftype is FileType.IMAGE:
        return _ingest_image(path)
    if ftype is FileType.DOCX:
        return _ingest_docx(path)
    raise UnsupportedFileError(
        f"Unsupported file type for '{path.name}'. Supported: PDF, JPG, PNG, DOCX."
    )


def _normalize_image_bytes(raw: bytes, max_dim: int) -> bytes:
    """Re-encode an image to PNG, applying EXIF rotation and downscaling."""
    with Image.open(io.BytesIO(raw)) as img:
        img = ImageOps.exif_transpose(img).convert("RGB")
        if max(img.size) > max_dim:
            img.thumbnail((max_dim, max_dim))
        out = io.BytesIO()
        img.save(out, format="PNG")
        return out.getvalue()


def render_page_png(
    doc: IngestedDocument,
    page_number: int,
    dpi: int | None = None,
    max_dim: int = 2200,
) -> bytes | None:
    """Rasterize one page (1-based) to PNG bytes, or ``None`` if not renderable.

    For PDFs the page is rendered at ``dpi``. For image files the (single) page
    is the image itself, re-encoded/normalized. DOCX has no renderable pages.
    """
    settings = get_settings()
    dpi = dpi or settings.pdf_render_dpi

    if doc.file_type is FileType.PDF:
        if page_number < 1 or page_number > doc.num_pages:
            return None
        pdf = fitz.open(doc.source_path)
        try:
            if pdf.needs_pass:
                pdf.authenticate("")
            zoom = dpi / 72.0
            pix = pdf[page_number - 1].get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            return pix.tobytes("png")
        finally:
            pdf.close()

    if doc.file_type is FileType.IMAGE and page_number == 1:
        return _normalize_image_bytes(doc.source_path.read_bytes(), max_dim)

    return None


def iter_page_images(
    doc: IngestedDocument,
    dpi: int | None = None,
    only_pages: list[int] | None = None,
) -> Iterator[tuple[int, bytes]]:
    """Yield ``(page_number, png_bytes)`` for renderable pages."""
    if not doc.can_render_images:
        return
    targets = only_pages or [p.page_number for p in doc.pages]
    for pn in targets:
        png = render_page_png(doc, pn, dpi=dpi)
        if png is not None:
            yield pn, png
